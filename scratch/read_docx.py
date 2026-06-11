import sys

try:
    import docx
    print("python-docx is installed")
    
    doc_path = "/Users/sebastianlilliecreutz/Downloads/Sales page copy.docx"
    doc = docx.Document(doc_path)
    print(f"Loaded document: {len(doc.paragraphs)} paragraphs")
    
    text = []
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text)
            
    # Check tables too
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text.append(cell.text)
                    
    full_text = "\n".join(text)
    
    # Write to temp file
    with open("scratch/docx_text.txt", "w", encoding="utf-8") as f:
        f.write(full_text)
    print("Wrote text to scratch/docx_text.txt")
    
    # Search for Maker Tour
    import re
    matches = re.findall(r'.{0,100}Maker Tour.{0,100}', full_text, re.IGNORECASE)
    print(f"Found {len(matches)} matches for Maker Tour:")
    for m in matches[:10]:
        print("-", m.strip())
        
except ImportError:
    print("python-docx is NOT installed")
except Exception as e:
    print("Error:", e)
